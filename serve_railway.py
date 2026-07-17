from __future__ import annotations

import base64
import io
import json
import os
import pathlib
import re
import urllib.error
import urllib.request
from copy import deepcopy
from datetime import date
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

ROOT = pathlib.Path(__file__).resolve().parent
TEMPLATE_DOCX = ROOT / 'source' / 'template-layout-word-terbaru.docx'
RETRYABLE_STATUS = {404, 429, 500, 502, 503, 504}
DEFAULT_PRIMARY_MODEL = 'gemini-2.5-flash'
DEFAULT_FALLBACK_MODELS = ['gemini-2.5-flash-lite', 'gemini-2.0-flash']


def normalize_model_list(*groups):
    result = []
    seen = set()
    for group in groups:
        if not group:
            continue
        if isinstance(group, str):
            group = [part.strip() for part in group.split(',')]
        for item in group:
            model = str(item or '').strip()
            if not model or model in seen:
                continue
            if any(c not in 'abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789._-' for c in model):
                continue
            seen.add(model)
            result.append(model)
    return result


def load_runtime_config() -> dict:
    primary = os.getenv('GEMINI_PRIMARY_MODEL', DEFAULT_PRIMARY_MODEL).strip() or DEFAULT_PRIMARY_MODEL
    fallback = normalize_model_list(os.getenv('GEMINI_FALLBACK_MODELS', ','.join(DEFAULT_FALLBACK_MODELS)))
    fallback = [m for m in fallback if m != primary]
    api_key = os.getenv('GEMINI_API_KEY', '').strip()
    return {
        'hasApiKey': bool(api_key),
        'apiKey': api_key,
        'primaryModel': primary,
        'fallbackModels': fallback,
        'allowBrowserOverride': False,
    }


def call_gemini(model: str, key: str, request_body: dict):
    data = json.dumps(request_body, separators=(',', ':'), ensure_ascii=False).encode('utf-8')
    req = urllib.request.Request(
        f'https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent',
        data=data,
        headers={'Content-Type': 'application/json', 'x-goog-api-key': key},
        method='POST',
    )
    try:
        with urllib.request.urlopen(req, timeout=240) as res:
            return res.status, res.read()
    except urllib.error.HTTPError as err:
        return err.code, err.read()


def _value(data: dict, key: str, default=''):
    value = data.get(key, default)
    if value is None:
        return default
    return value


def _parse_iso(value: str):
    if not isinstance(value, str) or not re.fullmatch(r'\d{4}-\d{2}-\d{2}', value):
        return None
    try:
        y, m, d = map(int, value.split('-'))
        return date(y, m, d)
    except ValueError:
        return None


def _date_jp(value: str, wide=False):
    parsed = _parse_iso(value)
    if not parsed:
        return ''
    if wide:
        return f'{parsed.year} 年　　 {parsed.month:02d} 月　　 {parsed.day:02d} 日'
    return f'{parsed.year} 年 {parsed.month:02d} 月 {parsed.day:02d} 日'


def _date_dash(value: str):
    parsed = _parse_iso(value)
    return f'{parsed.day:02d} - {parsed.month:02d} - {parsed.year}' if parsed else ''


def _date_id_wide(value: str):
    parsed = _parse_iso(value)
    return f'Tanggal　{parsed.day:02d}　Bulan　{parsed.month:02d}　Tahun {parsed.year}' if parsed else ''


def _age(birth: str, exam: str):
    b = _parse_iso(birth)
    e = _parse_iso(exam)
    if not b or not e:
        return ''
    years = e.year - b.year - ((e.month, e.day) < (b.month, b.day))
    return max(0, years)


def _bmi(height, weight):
    try:
        h = float(height) / 100.0
        w = float(weight)
        if h <= 0:
            return ''
        return f'{w / (h * h):.2f}'
    except (TypeError, ValueError, ZeroDivisionError):
        return ''


def _bilingual(indonesian, japanese, separator=' / '):
    left = str(indonesian or '').strip()
    right = str(japanese or '').strip()
    if left and right:
        return f'{left}{separator}{right}'
    return left or right


def _replace_paragraph_text(paragraph, text, *, size_pt=None, bold=None):
    text = str(text or '')
    runs = paragraph.runs
    if runs:
        target = runs[0]
        target.text = text
        for run in runs[1:]:
            run.text = ''
    else:
        target = paragraph.add_run(text)
    if size_pt is not None:
        target.font.size = Pt(size_pt)
    if bold is not None:
        target.bold = bool(bold)


def _set_cell_lines(cell, text, *, size_pt=None):
    lines = str(text or '').split('\n')
    paragraphs = cell.paragraphs
    if not paragraphs:
        paragraphs = [cell.add_paragraph()]
    for index, paragraph in enumerate(paragraphs):
        if index < len(lines):
            value = lines[index]
        elif index == len(paragraphs) - 1 and len(lines) > len(paragraphs):
            value = '\n'.join(lines[index:])
        else:
            value = ''
        _replace_paragraph_text(paragraph, value, size_pt=size_pt)
    if len(lines) > len(paragraphs):
        extra = '\n'.join(lines[len(paragraphs) - 1:])
        _replace_paragraph_text(paragraphs[-1], extra, size_pt=size_pt)


def _set_specific_paragraph(cell, index: int, text: str, *, size_pt=None):
    paragraphs = cell.paragraphs
    while len(paragraphs) <= index:
        cell.add_paragraph()
        paragraphs = cell.paragraphs
    _replace_paragraph_text(paragraphs[index], text, size_pt=size_pt)


def _set_exact_xml_text(container, old: str, new: str):
    for node in container._element.xpath('.//w:t'):
        if node.text == old:
            node.text = new


def _status_jp_id(value):
    positive = str(value or '').lower().startswith('pos')
    return 'ポジティブ\nPositif' if positive else 'ネガティブ\nNegatif'


def _vision(value, assisted):
    val = str(value or '').strip()
    if not val:
        return ''
    return f'（ {val} ）' if assisted else f'{val}　（　）'


def _hearing(value_1000, value_4000):
    def line(value):
        abnormal = str(value or '').lower().startswith('gang')
        if abnormal:
            return '1 所見なし　② 所見あり\n1 Normal　② Gangguan'
        return '① 所見なし　2 所見あり\n① Normal　2 Gangguan'
    return f'{line(value_1000)}\n{line(value_4000)}'


def _decode_stamp(value: str):
    if not value:
        return None
    if value.startswith('data:') and ',' in value:
        header, encoded = value.split(',', 1)
        if ';base64' not in header:
            return None
        try:
            return base64.b64decode(encoded, validate=False)
        except Exception:
            return None
    safe_path = pathlib.Path(value)
    if not safe_path.is_absolute():
        candidate = (ROOT / safe_path).resolve()
        try:
            candidate.relative_to(ROOT.resolve())
        except ValueError:
            return None
        if candidate.exists() and candidate.is_file():
            return candidate.read_bytes()
    return None


def build_word_document(data: dict) -> bytes:
    if not TEMPLATE_DOCX.exists():
        raise FileNotFoundError('Template Word tidak ditemukan di folder source.')

    document = Document(str(TEMPLATE_DOCX))
    tables = document.tables
    if len(tables) < 4:
        raise ValueError('Struktur template Word tidak valid.')

    nama = str(_value(data, 'nama')).strip()
    birth = str(_value(data, 'tglLahir')).strip()
    exam = str(_value(data, 'tglPeriksa')).strip()
    declaration = str(_value(data, 'tglDeklarasi')).strip()
    doc_date = str(_value(data, 'tglDokumen')).strip()
    age = _age(birth, exam)
    bmi = _bmi(_value(data, 'tinggi'), _value(data, 'berat'))

    # Halaman 1
    if len(document.paragraphs) > 5:
        _replace_paragraph_text(document.paragraphs[5], nama)
    declaration_table = tables[0]
    _set_cell_lines(declaration_table.rows[0].cells[1], _date_jp(declaration, wide=True), size_pt=10)
    _set_cell_lines(declaration_table.rows[1].cells[1], _date_id_wide(declaration), size_pt=10)

    # Halaman 2
    t = tables[1]
    _set_cell_lines(t.rows[0].cells[1], nama, size_pt=9)
    _set_cell_lines(t.rows[0].cells[4], f'{_date_jp(birth)}\n{_date_dash(birth)}', size_pt=8.5)
    _set_cell_lines(t.rows[0].cells[8], f'{_date_jp(exam)}\n{_date_dash(exam)}', size_pt=8.5)

    gender = str(_value(data, 'jenisKelamin', 'L')).upper()
    gender_text = '○男　・　女\nLaki-laki /\nPerempuan' if gender == 'L' else '男　・　○女\nLaki-laki /\nPerempuan'
    gender_cell = t.rows[1].cells[4]
    _set_cell_lines(gender_cell, gender_text, size_pt=8.5)
    for paragraph in gender_cell.paragraphs:
        for run in paragraph.runs:
            run.underline = False
            run.font.strike = False
    _set_cell_lines(t.rows[1].cells[8], f'{age}　歳\n{age}　tahun' if age != '' else '', size_pt=8.5)

    work = _bilingual(_value(data, 'riwayatKerjaId'), _value(data, 'riwayatKerjaJp'), '\n')
    history = _bilingual(_value(data, 'riwayatSakitId'), _value(data, 'riwayatSakitJp'), '\n/　')
    subjective = _bilingual(_value(data, 'gejalaSubId'), _value(data, 'gejalaSubJp'), '\n/　')
    objective = _bilingual(_value(data, 'gejalaObjId'), _value(data, 'gejalaObjJp'), '\n/　')
    _set_cell_lines(t.rows[2].cells[1], work, size_pt=8.5)
    _set_cell_lines(t.rows[5].cells[1], history, size_pt=8.5)
    _set_cell_lines(t.rows[8].cells[1], subjective, size_pt=8.5)
    _set_cell_lines(t.rows[11].cells[1], objective, size_pt=8.5)

    _set_cell_lines(t.rows[2].cells[9], f"{_value(data, 'tekananDarah')} mm/Hg" if _value(data, 'tekananDarah') else '', size_pt=8.5)
    _set_cell_lines(t.rows[3].cells[9], f"{_value(data, 'hb')} g/dℓ" if _value(data, 'hb') else '', size_pt=8.5)
    _set_cell_lines(t.rows[4].cells[9], f"{_value(data, 'rbc')} 万/mm³" if _value(data, 'rbc') else '', size_pt=8.5)
    _set_cell_lines(t.rows[5].cells[9], f"{_value(data, 'got')} IU/ℓ" if _value(data, 'got') else '', size_pt=8.5)
    _set_cell_lines(t.rows[6].cells[9], f"{_value(data, 'gpt')} IU/ℓ" if _value(data, 'gpt') else '', size_pt=8.5)
    _set_cell_lines(t.rows[7].cells[9], f"{_value(data, 'ggtp')} IU/ℓ" if _value(data, 'ggtp') else '', size_pt=8.5)
    _set_cell_lines(t.rows[8].cells[9], f"{_value(data, 'ldl')} mg/dℓ" if _value(data, 'ldl') else '', size_pt=8.5)
    _set_cell_lines(t.rows[9].cells[9], f"{_value(data, 'hdl')} mg/dℓ" if _value(data, 'hdl') else '', size_pt=8.5)
    _set_cell_lines(t.rows[10].cells[9], f"{_value(data, 'trigliserida')} mg/dℓ" if _value(data, 'trigliserida') else '', size_pt=8.5)
    glucose_prefix = '*' if bool(_value(data, 'gulaDarahBintang', False)) else ''
    glucose_value = _value(data, 'gulaDarah')
    _set_cell_lines(t.rows[11].cells[9], f'{glucose_prefix}{glucose_value} mg/dℓ' if glucose_value else '', size_pt=8.5)
    _set_cell_lines(t.rows[12].cells[9], _status_jp_id(_value(data, 'glukosaUrine')), size_pt=8)
    _set_cell_lines(t.rows[13].cells[9], _status_jp_id(_value(data, 'proteinUrine')), size_pt=8)
    _set_cell_lines(t.rows[14].cells[1], f"{_value(data, 'tinggi')} cm" if _value(data, 'tinggi') else '', size_pt=8.5)
    _set_cell_lines(t.rows[15].cells[1], f"{_value(data, 'berat')} kg" if _value(data, 'berat') else '', size_pt=8.5)

    ekg = _bilingual(_value(data, 'ekgId'), _value(data, 'ekgJp'), ' / ')
    other = _bilingual(_value(data, 'pemeriksaanLainId'), _value(data, 'pemeriksaanLainJp'), ' / ')
    _set_cell_lines(t.rows[15].cells[6], ekg, size_pt=8.5)
    _set_cell_lines(t.rows[16].cells[6], other, size_pt=8.5)

    # Halaman 3
    t3 = tables[2]
    _set_cell_lines(t3.rows[0].cells[2], f'{bmi} kg/m²' if bmi else '', size_pt=8.5)
    _set_cell_lines(t3.rows[1].cells[2], f"{_value(data, 'lingkarPerut')} cm" if _value(data, 'lingkarPerut') else '', size_pt=8.5)
    assisted = bool(_value(data, 'alatBantuMata', False))
    _set_cell_lines(t3.rows[2].cells[2], _vision(_value(data, 'mataKanan'), assisted), size_pt=8.5)
    _set_cell_lines(t3.rows[3].cells[2], _vision(_value(data, 'mataKiri'), assisted), size_pt=8.5)
    _set_cell_lines(t3.rows[4].cells[2], _hearing(_value(data, 'telingaKanan1000'), _value(data, 'telingaKanan4000')), size_pt=7)
    _set_cell_lines(t3.rows[5].cells[2], _hearing(_value(data, 'telingaKiri1000'), _value(data, 'telingaKiri4000')), size_pt=7)

    direct = str(_value(data, 'rontgenMetode', 'Langsung')) != 'Tidak langsung'
    method_jp = '○直接　　　　　間接' if direct else '直接　　　　　○間接'
    method_id = '○Langsung　　 Tidak langsung' if direct else 'Langsung　　 ○Tidak langsung'
    xray = (
        f'{method_jp}\n{method_id}\n'
        f'撮影　　 {_date_jp(_value(data, "rontgenTanggal"))}\n'
        f'Diambil tanggal　{_date_dash(_value(data, "rontgenTanggal"))}\n'
        f'No.　{_value(data, "rontgenNo")}\n'
        f'所見: {_value(data, "rontgenTemuanJp")}\n'
        f'Temuan: {_value(data, "rontgenTemuanId")}'
    )
    _set_cell_lines(t3.rows[7].cells[2], xray, size_pt=7)

    diagnosis = _bilingual(_value(data, 'diagnosisId'), _value(data, 'diagnosisJp'), ' / ')
    diagnosis_cell = t3.rows[1].cells[3]
    _set_specific_paragraph(diagnosis_cell, 1, diagnosis, size_pt=8)
    fit = str(_value(data, 'fitStatus', 'FIT')).upper()
    _set_exact_xml_text(diagnosis_cell, 'FIT', '● FIT' if fit == 'FIT' else 'FIT')
    _set_exact_xml_text(diagnosis_cell, 'UNFIT', '● UNFIT' if fit == 'UNFIT' else 'UNFIT')
    _set_cell_lines(t3.rows[8].cells[3], _bilingual(_value(data, 'keteranganId'), _value(data, 'keteranganJp'), ' / '), size_pt=8)

    footer = tables[3]
    _set_cell_lines(footer.rows[0].cells[0], f'作成年月日　　　　{_date_jp(doc_date)}\nTanggal pembuatan:　{_date_id_wide(doc_date)}', size_pt=8)

    stamp_bytes = _decode_stamp(str(_value(data, 'stampData', '')))
    doctor_name = str(_value(data, 'dokterNama')).strip()
    clinic_name = str(_value(data, 'klinikNama')).strip()
    stamp_cell = footer.rows[0].cells[1]
    stamp_paragraph = stamp_cell.paragraphs[0]
    if stamp_bytes:
        # Clear any existing text in the target paragraph, then insert the uploaded stamp.
        _replace_paragraph_text(stamp_paragraph, '')
        stamp_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = stamp_paragraph.add_run()
        run.add_picture(io.BytesIO(stamp_bytes), width=Mm(45))
        if doctor_name:
            name_paragraph = stamp_cell.add_paragraph()
            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_paragraph.add_run(f'({doctor_name})')
            name_run.font.size = Pt(7.5)
    else:
        lines = [line for line in (clinic_name, f'({doctor_name})' if doctor_name else '') if line]
        _set_cell_lines(stamp_cell, '\n'.join(lines), size_pt=8)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


class Handler(SimpleHTTPRequestHandler):
    protocol_version = 'HTTP/1.1'

    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(ROOT), **kwargs)

    def do_GET(self):
        if self.path in ('/api/gemini-config', '/api/gemini-config/'):
            cfg = load_runtime_config()
            body = json.dumps({
                'hasApiKey': cfg['hasApiKey'],
                'primaryModel': cfg['primaryModel'],
                'fallbackModels': cfg['fallbackModels'],
                'allowBrowserOverride': False,
            }).encode('utf-8')
            self.send_response(200)
            self.send_header('Content-Type', 'application/json; charset=utf-8')
            self.send_header('Content-Length', str(len(body)))
            self.end_headers()
            self.wfile.write(body)
            return
        sensitive = ('gemini.config.json', 'gemini.config.local.json', '.env')
        if any(self.path.endswith(name) for name in sensitive):
            self.send_error(404)
            return
        return super().do_GET()

    def do_POST(self):
        if self.path == '/api/export-word':
            try:
                length = int(self.headers.get('Content-Length', '0'))
                if length <= 0 or length > 12 * 1024 * 1024:
                    raise ValueError('Ukuran data export Word tidak valid.')
                payload = json.loads(self.rfile.read(length).decode('utf-8'))
                data = payload.get('data')
                if not isinstance(data, dict):
                    raise ValueError('Data formulir Word tidak valid.')
                raw_name = str(payload.get('filename') or 'MCU_Medical_Checkup.docx')
                filename = re.sub(r'[^A-Za-z0-9._-]+', '_', raw_name).strip('._') or 'MCU_Medical_Checkup.docx'
                if not filename.lower().endswith('.docx'):
                    filename += '.docx'
                body = build_word_document(data)
                self.send_response(200)
                self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
                self.send_header('Content-Length', str(len(body)))
                self.end_headers()
                self.wfile.write(body)
            except Exception as exc:
                body = json.dumps({'error': {'message': str(exc)}}).encode('utf-8')
                self.send_response(400)
                self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.send_header('Content-Length', str(len(body)))
                self.end_headers()
                self.wfile.write(body)
            return

        if self.path != '/api/gemini':
            self.send_error(404)
            return
        try:
            length = int(self.headers.get('Content-Length', '0'))
            if length <= 0 or length > 80 * 1024 * 1024:
                raise ValueError('Ukuran permintaan tidak valid.')
            payload = json.loads(self.rfile.read(length).decode('utf-8'))
            request_body = payload.get('request')
            if not isinstance(request_body, dict):
                raise ValueError('Body permintaan Gemini tidak valid.')

            cfg = load_runtime_config()
            key = cfg['apiKey']
            if not key:
                raise ValueError('GEMINI_API_KEY belum diatur pada environment variables.')

            models = normalize_model_list(cfg['primaryModel'], cfg['fallbackModels'])
            if not models:
                raise ValueError('Model Gemini belum diatur.')

            attempts = []
            final_status, final_body = 500, b''
            for idx, model in enumerate(models):
                status, body = call_gemini(model, key, request_body)
                attempts.append({'model': model, 'status': status})
                if 200 <= status < 300:
                    self.send_response(status)
                    self.send_header('Content-Type', 'application/json; charset=utf-8')
                    self.send_header('Content-Length', str(len(body)))
                    self.send_header('X-Gemini-Model-Used', model)
                    self.send_header('X-Gemini-Attempts', json.dumps(attempts, ensure_ascii=False))
                    self.end_headers()
                    self.wfile.write(body)
                    return
                final_status, final_body = status, body
                if status in (400, 401, 403):
                    break
                if status not in RETRYABLE_STATUS and idx == len(models) - 1:
                    break

            self.send_response(final_status)
            self.send_header('Content-Type', 'application/json; charset=utf-8')
            self.send_header('Content-Length', str(len(final_body)))
            self.send_header('X-Gemini-Attempts', json.dumps(attempts, ensure_ascii=False))
            self.end_headers()
            self.wfile.write(final_body)
        except Exception as exc:
            body = json.dumps({'error': {'message': str(exc)}}).encode('utf-8')
            self.send_response(400)
            self.send_header('Content-Type', 'application/json; charset=utf-8')
            self.send_header('Content-Length', str(len(body)))
            self.end_headers()
            self.wfile.write(body)


if __name__ == '__main__':
    host = os.getenv('HOST', '0.0.0.0')
    port = int(os.getenv('PORT', '8080'))
    server = ThreadingHTTPServer((host, port), Handler)
    print(f'Generator MCU listening on http://{host}:{port}')
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        pass
    finally:
        server.server_close()
